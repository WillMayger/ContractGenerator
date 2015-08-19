from docx import Document
from datetime import datetime

# ToDo
# fix so it will use the GBP sign symbol in method compensation

class ContractGenerator:

    def __init__(self, client_name, client_address, client_town, client_country, client_postcode, client_fax,
                 client_email,
                 service_provider_name, service_provider_address, service_provider_town,
                 service_provider_country, service_provider_postcode,
                 service_provider_fax, service_provider_email, service_provided, currency,
                 cost):

        # date time
        now = datetime.now()
        if 4 <= now.day.__int__() <= 20 or 24 <= now.day.__int__() <= 30:
            suffix = "th"
        else:
            suffix = ["st", "nd", "rd"][now.day.__int__() % 10 - 1]

        self.date = now.day.__str__() + suffix
        self.month = now.month.__str__()
        self.year = now.year.__str__()

        # client info
        self.client_name = client_name
        self.client_address = client_address
        self.client_town = client_town
        self.client_country = client_country
        self.client_postcode = client_postcode
        self.client_fax = client_fax
        self.client_email = client_email

        # currency
        self.currency_unit = currency

        # cost
        self.cost = cost

        # service provider info
        self.service_provider_name = service_provider_name
        self.service_provider_address = service_provider_address
        self.service_provider_town = service_provider_town
        self.service_provider_country = service_provider_country
        self.service_provider_postcode = service_provider_postcode
        self.service_provided = service_provided
        self.service_provider_fax = service_provider_fax
        self.service_provider_email = service_provider_email

        # new instance of Document
        self.contract = Document()

    def heading(self):
        self.contract.add_heading('GENERAL SERVICE AGREEMENT', 0)
        self.contract.add_heading(
            'THIS GENERAL SERVICE AGREEMENT (the "Agreement") dated this {0} day of {1}, {2}'.format(
                self.date, self.month, self.year
            ), 1)

    # first part of the contract document adding the addresses
    def client_and_service_info(self):

        self.contract.add_heading(
            'BETWEEN', 2)

        self.contract.add_paragraph(
            '{0} of {1}, {2}, {3},\n {4} \n (the "Customer")'.format(
                self.client_name, self.client_address, self.client_town, self.client_country,
                self.client_postcode
            ))

        self.contract.add_heading(
            '- AND -', 2)

        self.contract.add_paragraph(
            '{0} of {1}, {2}, {3},\n {4} \n (the "Service Provider")'.format(
                self.service_provider_name, self.service_provider_address,
                self.service_provider_town, self.service_provider_country,
                self.service_provider_postcode
            ))

    def background(self):
        self.contract.add_heading(
            'BACKGROUND:', 2)

        self.contract.add_paragraph(
            '1. The Customer is of the opinion that the Service Provider has the necessary qualifications,' +
            ' experience and abilities to provide services to the Customer.'
            )
        self.contract.add_paragraph(
            '2. The Service Provider is agreeable to providing such services to the Customer on the terms' +
            ' and conditions set out in this Agreement.'
            )

    def in_consideration_of(self):

        self.contract.add_heading(
            'IN CONSIDERATION OF:', 2)

        self.contract.add_paragraph(
            'IN CONSIDERATION OF the matters described above and of the mutual benefits and obligations set' +
            ' forth in this Agreement, the receipt and sufficiency of which consideration is hereby acknowledged,' +
            ' the Customer and the Service Provider (individually the "Party" and collectively the "Parties" to ' +
            'this Agreement) agree as follows:'
            )

        self.contract.add_heading(
            '1. Services Provided\n',
            6)

        self.contract.add_paragraph(
            '2. The Customer hereby agrees to engage the Service Provider to provide the Customer with services'
            ' (the "Services") consisting of:'
            )

        self.contract.add_paragraph(
            '  2.1 ' + self.service_provided
        )

        self.contract.add_paragraph(
            '3. The Services will also include any other tasks which the Parties may agree on.' +
            ' The Service Provider hereby agrees to provide such Services to the Customer.'
        )

    def terms_of_agreement(self):

        self.contract.add_heading(
            'Term of Agreement:\n', 3)

        self.contract.add_paragraph(
            '1. The term of this Agreement (the "Term") will begin on the date of this Agreement and will remain ' +
            'in full force and effect until the completion of the Services, subject to earlier termination as ' +
            'provided in this Agreement. The Term of this Agreement may be extended by mutual written agreement ' +
            'of the Parties.'
            )

        self.contract.add_paragraph(
            '2. In the event that either Party wishes to terminate this Agreement, that Party will be required ' +
            'to provide 3 days notice to the other Party.'
            )

    def performance(self):

        self.contract.add_heading(
            'Performance:\n', 3)

        self.contract.add_paragraph(
            '1. The Parties agree to do everything necessary to ensure that the terms of this Agreement take effect.'
            )

    def currency(self):

        self.contract.add_heading(
            'Currency:\n', 3
        )

        self.contract.add_paragraph(
            '1. Except as otherwise provided in this Agreement, all monetary amounts referred' +
            ' to in this Agreement are in {0}.'.format(self.currency_unit)
        )

    def compensation(self):

        self.contract.add_heading(
            'Compensation:\n', 3
        )

        self.contract.add_paragraph(
            '1. For the services rendered by the Service Provider as required by this Agreement, ' +
            'the Customer will provide compensation (the "Compensation") to the Service Provider ' +
            'of a fixed amount of ' +
             u"\xA3" +
            ' {0}.'.format(self.cost)
        )

        self.contract.add_paragraph(
            '2. The Compensation will be payable upon completion of the Services.'
        )

        self.contract.add_paragraph(
            '3. The Service Provider will be responsible for all income tax liabilities' +
            ' and National Insurance or similar contributions relating to the Compensation' +
            ' and the Service Provider will indemnify the Company in respect of any such payments' +
            ' required to be made by the Company.'
        )

    def confidentiality(self):

        self.contract.add_heading(
            'Confidentiality:\n', 3
        )

        self.contract.add_paragraph(
            '1. Confidential information (the "Confidential Information") refers to any ' +
            'data or information relating to the Customer, whether business or personal, ' +
            'which would reasonably be considered to be private or proprietary to the ' +
            'Customer and that is not generally known and where the release of that ' +
            'Confidential Information could reasonably be expected to cause harm to the Customer.'
        )

        self.contract.add_paragraph(
            '2. The Service Provider agrees that they will not disclose, divulge, reveal, ' +
            'report or use, for any purpose, any Confidential Information which the Service ' +
            'Provider has obtained, except as authorized by the Customer. This obligation ' +
            'will survive indefinitely upon termination of this Agreement.'
        )

        self.contract.add_paragraph(
            '3. All written and oral information and material disclosed or provided by the ' +
            'Customer to the Service Provider under this Agreement is Confidential Information ' +
            'regardless of whether it was provided before or after the date of this Agreement ' +
            'or how it was provided to the Service Provider.'
        )

    def ownership_materials_intellectual_property(self):

        self.contract.add_heading(
            'Ownership of Materials and Intellectual Property:\n',
            3
        )

        self.contract.add_paragraph(
            '1. All intellectual property and related material (the "Intellectual Property") ' +
            'including any related work in progress that is developed or produced under this ' +
            'Agreement, will be the sole property of the Customer. The use of the Intellectual ' +
            'Property by the Customer will not be restricted in any manner.'
        )

        self.contract.add_paragraph(
            '2. The Service Provider may not use the Intellectual Property for any ' +
            'purpose other than that contracted for in this Agreement except with the ' +
            'written consent of the Customer. The Service Provider will be responsible ' +
            'for any and all damages resulting from the unauthorized use of the Intellectual Property.'
        )

    def return_of_property(self):

        self.contract.add_heading(
            'Return of Property:\n', 3
        )

        self.contract.add_paragraph(
            '1. Upon the expiry or termination of this Agreement, the Service Provider ' +
            'will return to the Customer any property, documentation, records, or Confidential ' +
            'Information which is the property of the Customer.'
        )

    def capacity_independent_contractor(self):

        self.contract.add_heading(
            'Capacity/Independent Contractor:\n',
            3
        )

        self.contract.add_paragraph(
            'In providing the Services under this Agreement it is ' +
            'expressly agreed that the Service Provider is acting as an ' +
            'independent contractor and not as an employee. The Service ' +
            'Provider and the Customer acknowledge that this Agreement does not create a ' +
            'partnership or joint venture between them, and is exclusively a contract for service.'
        )

    def notice(self):

        self.contract.add_heading(
            'Notice:\n', 3
        )

        self.contract.add_paragraph(
            'All notices, requests, demands or other communications required or ' +
            'permitted by the terms of this Agreement will be given in writing and ' +
            'delivered to the Parties of this Agreement as follows:'
        )

        self.contract.add_paragraph(
            str(
                '1. {0}\n' +
                '2. {1}\n' +
                '3. {2}, {3}, {4}\n' +
                '4. Fax: {5}\n' +
                '5. Email: {6}\n'
            ).format(
            self.client_name,
            self.client_address,
            self.client_town, self.client_country, self.client_postcode,
            self.client_fax,
            self.client_email
            )
        )

        self.contract.add_paragraph(
            str(
                '1. {0}\n' +
                '2. {1}\n' +
                '3. {2}, {3}, {4}\n' +
                '4. Fax: {5}\n' +
                '5. Email: {6}\n'
            ).format(
            self.service_provider_name,
            self.service_provider_address,
            self.service_provider_town, self.service_provider_country, self.service_provider_postcode,
            self.service_provider_fax,
            self.service_provider_email
            )
        )

        self.contract.add_paragraph(
            'or to such other address as any Party may from time to time notify the other.'
        )

    def indemnification(self):

        self.contract.add_heading(
            'Indemnification:\n', 3
        )

        self.contract.add_paragraph(
            '1. Each Party to this Agreement will indemnify and hold harmless the other ' +
            'Party, as permitted by law, from and against any and all claims, losses, damages, ' +
            'liabilities, penalties, punitive damages, expenses, reasonable legal fees and costs ' +
            'of any kind or amount whatsoever to the extent that any of the foregoing is directly ' +
            'or proximately caused by the negligent or wilful acts or omissions of the indemnifying ' +
            'Party or its agents or representatives and which result from or arise out of the ' +
            'indemnifying Party\'s participation in this Agreement. This indemnification will ' +
            'survive the termination of this Agreement.'
        )

    def modification_of_agreement(self):

        self.contract.add_heading(
            'Modification of Agreement:\n', 3
        )

        self.contract.add_paragraph(
            '1. Any amendment or modification of this Agreement or additional obligation assumed ' +
            'by either Party in connection with this Agreement will only be binding if evidenced ' +
            'in writing signed by each Party or an authorized representative of each Party.'
        )

    def time_of_the_essence(self):

        self.contract.add_heading(
            'Time of the Essence:\n', 3
        )

        self.contract.add_paragraph(
            '1. Time is of the essence in this Agreement. No extension or variation of ' +
            'this Agreement will operate as a waiver of this provision.'
        )

    def assignment(self):

        self.contract.add_heading(
            'Assignment:\n', 3
        )

        self.contract.add_paragraph(
            '1. The Service Provider will not voluntarily or by operation of law assign or ' +
            'otherwise transfer its obligations under this Agreement without the prior ' +
            'written consent of the Customer.'
        )

    def entire_agreement(self):
        self.contract.add_heading(
            'Entire Agreement:\n', 3
        )

        self.contract.add_paragraph(
            '1. It is agreed that there is no representation, warranty, collateral agreement ' +
            'or condition affecting this Agreement except as expressly provided in this Agreement.'
        )

    def enurement(self):

        self.contract.add_heading(
            'Enurement:\n', 3
        )

        self.contract.add_paragraph(
            '1. This Agreement will enure to the benefit of and be binding on the Parties ' +
            'and their respective heirs, executors, administrators, successors and permitted assigns.'
        )

    def titles_headings(self):

        self.contract.add_heading(
            'Titles/Headings:\n', 3
        )

        self.contract.add_paragraph(
            '1. Headings are inserted for the convenience of the Parties ' +
            'only and are not to be considered when interpreting this Agreement.'
        )

    def gender(self):

        self.contract.add_heading(
            'Gender:\n', 3
        )

        self.contract.add_paragraph(
            '1. Words in the singular mean and include the plural and vice versa. ' +
            'Words in the masculine mean and include the feminine and vice versa.'
        )

    def governing_law(self):

        self.contract.add_heading(
            'Governing Law:\n', 3
        )

        self.contract.add_paragraph(
            '1. It is the intention of the Parties to this Agreement that this Agreement ' +
            'and the performance under this Agreement, and all suits and special proceedings ' +
            'under this Agreement, be construed in accordance with and governed, to the ' +
            'exclusion of the law of any other forum, by the laws of the Country of England, ' +
            'without regard to the jurisdiction in which any action or special proceeding may be instituted.'
        )

    def severability(self):

        self.contract.add_heading(
            'Severability:\n', 3
        )

        self.contract.add_paragraph(
            'In the event that any of the provisions of this Agreement are held to be ' +
            'invalid or unenforceable in whole or in part, all other provisions will ' +
            'nevertheless continue to be valid and enforceable with the invalid or unenforceable ' +
            'parts severed from the remainder of this Agreement.'
        )

    def waiver(self):

        self.contract.add_heading(
            'Waiver:\n', 3
        )

        self.contract.add_paragraph(
            'The waiver by either Party of a breach, default, delay or omission of any of the ' +
            'provisions of this Agreement by the other Party will not be construed as a waiver ' +
            'of any subsequent breach of the same or other provisions.'
        )

    def in_witness_of(self):

        self.contract.add_heading(
            'IN WITNESS WHEREOF:\n', 3
        )

        self.contract.add_paragraph(
            'IN WITNESS WHEREOF the Parties have duly affixed their signatures under hand ' +
            'and seal on this {0} day of {1}, {2}.\n'.format(
                self.date, self.month, self.year
            )
        )

    def signatures(self):

        self.contract.add_paragraph(
            '...................................................................................\n' +
            '{0}\n'.format(self.client_name)
        )

        self.contract.add_paragraph(
            '...................................................................................\n' +
            '{0}\n'.format(self.service_provider_name)
        )


    def save(self):

        self.contract.save('contract.docx')


cg = ContractGenerator("Mr Example", "93 london drive", "camden", "london", "LON DON", "fax", "example@example.com",
                       "Mr Company Example", "93 Company drive", "Regents Place", "London", "LON DON", "fax",
                       "company@company.com",
                       "Responsive Website Development with up to date responsive devices to work across all devices.",
                       "GBP", "100000.00")

cg.heading()
cg.client_and_service_info()
cg.background()
cg.in_consideration_of()
cg.terms_of_agreement()
cg.performance()
cg.currency()
cg.compensation()
cg.confidentiality()
cg.ownership_materials_intellectual_property()
cg.return_of_property()
cg.capacity_independent_contractor()
cg.notice()
cg.indemnification()
cg.modification_of_agreement()
cg.time_of_the_essence()
cg.assignment()
cg.entire_agreement()
cg.enurement()
cg.titles_headings()
cg.gender()
cg.governing_law()
cg.severability()
cg.waiver()
cg.in_witness_of()
cg.signatures()
cg.save()
